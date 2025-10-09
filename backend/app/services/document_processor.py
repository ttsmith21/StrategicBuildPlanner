"""
Document Processing Service
Handles PDF, DOCX, and TXT file processing
"""

import io
import logging
from typing import BinaryIO, Dict, List, Tuple
from pathlib import Path
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document types"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    MAX_FILE_SIZE_MB = 50
    
    @staticmethod
    def validate_file(filename: str, file_size: int) -> Tuple[bool, str]:
        """
        Validate file type and size
        
        Args:
            filename: Original filename
            file_size: File size in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check extension
        ext = Path(filename).suffix.lower()
        if ext not in DocumentProcessor.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type: {ext}. Supported: {', '.join(DocumentProcessor.SUPPORTED_EXTENSIONS)}"
        
        # Check size
        max_size_bytes = DocumentProcessor.MAX_FILE_SIZE_MB * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"File too large: {file_size / 1024 / 1024:.1f}MB. Max: {DocumentProcessor.MAX_FILE_SIZE_MB}MB"
        
        return True, ""
    
    @staticmethod
    async def extract_text_from_pdf(file: BinaryIO, filename: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file: File object (binary mode)
            filename: Original filename (for logging)
            
        Returns:
            Extracted text
        """
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
            
            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from PDF: {filename} ({len(pdf_reader.pages)} pages)")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {filename}: {str(e)}")
            raise ValueError(f"Could not process PDF file: {str(e)}")
    
    @staticmethod
    async def extract_text_from_docx(file: BinaryIO, filename: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file: File object (binary mode)
            filename: Original filename (for logging)
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(file)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX: {filename}")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {filename}: {str(e)}")
            raise ValueError(f"Could not process DOCX file: {str(e)}")
    
    @staticmethod
    async def extract_text_from_txt(file: BinaryIO, filename: str) -> str:
        """
        Extract text from TXT file
        
        Args:
            file: File object (binary mode)
            filename: Original filename (for logging)
            
        Returns:
            Extracted text
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            content = file.read()
            
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed for {filename}, trying latin-1")
                text = content.decode('latin-1')
            
            logger.info(f"Extracted {len(text)} characters from TXT: {filename}")
            return text
            
        except Exception as e:
            logger.error(f"Failed to extract text from TXT {filename}: {str(e)}")
            raise ValueError(f"Could not process TXT file: {str(e)}")
    
    @staticmethod
    async def process_file(file: BinaryIO, filename: str) -> Dict[str, str]:
        """
        Process a file and extract text based on type
        
        Args:
            file: File object (binary mode)
            filename: Original filename
            
        Returns:
            Dictionary with extracted text and metadata
        """
        ext = Path(filename).suffix.lower()
        
        # Reset file pointer
        file.seek(0)
        
        # Extract text based on file type
        if ext == '.pdf':
            text = await DocumentProcessor.extract_text_from_pdf(file, filename)
        elif ext == '.docx':
            text = await DocumentProcessor.extract_text_from_docx(file, filename)
        elif ext == '.txt':
            text = await DocumentProcessor.extract_text_from_txt(file, filename)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Reset file pointer again for upload
        file.seek(0)
        
        return {
            "filename": filename,
            "extension": ext,
            "text": text,
            "char_count": len(text),
            "word_count": len(text.split())
        }
    
    @staticmethod
    async def process_multiple_files(
        files: List[Tuple[BinaryIO, str]]
    ) -> List[Dict[str, str]]:
        """
        Process multiple files
        
        Args:
            files: List of (file, filename) tuples
            
        Returns:
            List of processed file metadata
        """
        results = []
        
        for file, filename in files:
            try:
                result = await DocumentProcessor.process_file(file, filename)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {filename}: {str(e)}")
                # Continue with other files
                results.append({
                    "filename": filename,
                    "error": str(e),
                    "char_count": 0,
                    "word_count": 0
                })
        
        return results
